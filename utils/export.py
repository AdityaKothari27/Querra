from fpdf import FPDF
from docx import Document
import os
from datetime import datetime

class ReportExporter:
    def __init__(self):
        self.export_dir = "exports"
        os.makedirs(self.export_dir, exist_ok=True)

    def _get_filename(self, query, format):
        """Generate a filename based on the query and timestamp"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_query = "".join(c if c.isalnum() else "_" for c in query)[:30]
        return f"{safe_query}_{timestamp}.{format.lower()}"

    def export_pdf(self, query, content):
        """Export report as PDF"""
        try:
            pdf = FPDF()
            pdf.add_page()
            
            # Set font
            pdf.set_font("Arial", "B", 16)
            pdf.cell(0, 10, "Research Report", ln=True, align="C")
            
            pdf.set_font("Arial", "B", 12)
            pdf.cell(0, 10, f"Topic: {query}", ln=True)
            
            pdf.set_font("Arial", "", 12)
            
            # Split content into lines and add to PDF
            pdf.multi_cell(0, 10, content)
            
            filename = self._get_filename(query, "pdf")
            filepath = os.path.join(self.export_dir, filename)
            pdf.output(filepath)
            
            return filepath
        except Exception as e:
            print(f"Error exporting to PDF: {e}")
            return None

    def export_docx(self, query, content):
        """Export report as DOCX"""
        try:
            doc = Document()
            
            # Add title
            doc.add_heading("Research Report", 0)
            
            # Add topic
            doc.add_heading(f"Topic: {query}", 1)
            
            # Add content
            doc.add_paragraph(content)
            
            filename = self._get_filename(query, "docx")
            filepath = os.path.join(self.export_dir, filename)
            doc.save(filepath)
            
            return filepath
        except Exception as e:
            print(f"Error exporting to DOCX: {e}")
            return None

    def export_txt(self, query, content):
        """Export report as TXT"""
        try:
            filename = self._get_filename(query, "txt")
            filepath = os.path.join(self.export_dir, filename)
            
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(f"Research Report\n\n")
                f.write(f"Topic: {query}\n\n")
                f.write(content)
            
            return filepath
        except Exception as e:
            print(f"Error exporting to TXT: {e}")
            return None